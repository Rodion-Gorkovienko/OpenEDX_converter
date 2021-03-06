import docx
import xml_elements
import xml_library
import xml_checkbox
import xml_multiple_choice
import xml_text_match
import xml_math_expression
import xml_numeric
from enum import Enum

class question_type(Enum):
    SINGLE_CHOICE = 0
    MULTIPLE_CHOICE = 1
    NUMERIC = 2
    TEXT_MATCH = 3
    MATH_EXPRESSION = 4

class shuffle(Enum):
    NO_SHUFFLE = 5
    SHUFFLE = 6

class partial_credit(Enum):
    NO_PARTIAL_CREDIT = 7
    PARTIAL_CREDIT = 8

class IncorrectSyntax(Exception):
    pass

def MathJax_border_fix(question):
    i = 0
    while (i != -1 and i < len(question)):
        start_tag = question.find("$$", i)
        if start_tag != -1:
            end_tag = question.find("$$", start_tag + 2)
            if end_tag == -1:
                raise IncorrectSyntax("An odd number of \"$$\" in the question")
            if question.find('\n', start_tag, end_tag) !=-1:
                raise IncorrectSyntax("Line feed inside MathJax formula")
            question = question[: start_tag] + "\\(" + question[start_tag + 2: end_tag] + "\\)" + question[end_tag + 2 :]
        i = start_tag
    return question

def create_checkbox_question(options, question):
    if question.count('\n*') == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    problem = xml_elements.problem()
    choiceresponse = xml_checkbox.choiceresponse(2, question)
    if options[2] == partial_credit.PARTIAL_CREDIT:
        choiceresponse.add_partial_credit()
    problem.add_content(choiceresponse)
    return problem

def create_multiple_choice_question(options, question):
    count = question.count('\n*')
    if count == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    if count > 1:
        raise IncorrectSyntax("Multiple correct answers are not allowed for this type of question")
    shuffle_set = False
    if options[1] == shuffle.SHUFFLE:
        shuffle_set = True
    problem = xml_elements.problem()
    multiple_choice_response = xml_multiple_choice.multiplechoiceresponse(2, question, shuffle_set)
    problem.add_content(multiple_choice_response)
    return problem

def create_text_match_question(question):
    if question.count('\n*') == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    problem = xml_elements.problem()
    stringresponse = xml_text_match.stringresponse(2, question)
    problem.add_content(stringresponse)
    return problem

def create_math_expression_question(question):
    if question.count('\n*') == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    problem = xml_elements.problem()
    stringresponse = xml_math_expression.formularesponse(2, question)
    problem.add_content(stringresponse)
    return problem

def create_numeric_question(question):
    if question.count('\n*') == 0:
        raise IncorrectSyntax("Question does not have a correct answer")
    problem = xml_elements.problem()
    numericalresponse = xml_numeric.numericalresponse(2, question)
    problem.add_content(numericalresponse)
    return problem

def parse_file(file_full_ref):
    doc = docx.Document(file_full_ref)

    #print(len(doc.paragraphs))

    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    united_text = str('\n'.join(text))
    #print(united_text)

    #Library creation
    lib = xml_library.library()
    properties_order = ["display_name", "library", "org"]
    i = 0
    i_prop = 0;
    if i != united_text.find("Question 1"):
        if united_text[i + 1] != "\n":
            lib.add_property(properties_order[i_prop], united_text[i: united_text.find("\n")])
            i_prop += 1
            i = united_text.find("\n")
        else:
            i += 1
        while(i != united_text.find("\nQuestion 1")):
            #print(i)
            #print(united_text.find("\nQuestion 1"))
            if united_text[i + 1] != "\n":
                lib.add_property(properties_order[i_prop], united_text[i + 1: united_text.find("\n", i + 1)])
                i_prop += 1
                i = united_text.find("\n", i + 1)
            else:
                i += 1

    #Questions creation
    problems = list()
    q_i = 1
    try:
        while(i != -1 and i < len(united_text)):
            if united_text[i :i + 10 + len(str(q_i))] != "\nQuestion " + str(q_i):
                raise IncorrectSyntax("Incorrect syntax of the question entry")
            options = [question_type.MULTIPLE_CHOICE, shuffle.NO_SHUFFLE, partial_credit.NO_PARTIAL_CREDIT]
            start_of_question = united_text.find("\n", i + 1)
            #start_of_answers = min(united_text.find("\n*"), united_text.find("\nA:"))
            end_of_question = united_text.find("\nQuestion", start_of_question)
            if end_of_question == -1:
                end_of_question = len(united_text)
            question_announcement = united_text[i + 1 : start_of_question]
            full_question = united_text[start_of_question : end_of_question]
            #print(question_announcement)
            #print(full_question)

            #Type recognizing
            if question_announcement.find("multiple choice") != -1:
                if full_question.count('\n*') > 1:
                    options[0] = question_type.MULTIPLE_CHOICE
                else:
                    options[0] = question_type.SINGLE_CHOICE
            if question_announcement.find("single correct answer") != -1:
                options[0] = question_type.SINGLE_CHOICE
            if question_announcement.find("multiple correct answers") != -1:
                options[0] = question_type.MULTIPLE_CHOICE
            if question_announcement.find("checkbox") != -1:
                options[0] = question_type.MULTIPLE_CHOICE
            if question_announcement.find("text match") != -1:
                options[0] = question_type.TEXT_MATCH
            if question_announcement.find("math expression") != -1:
                options[0] = question_type.MATH_EXPRESSION
            if question_announcement.find("numeric") != -1:
                options[0] = question_type.NUMERIC
            if question_announcement.find("regex") != -1:
                raise IncorrectSyntax("Unsupported question type: regular expression")
            if question_announcement.find("regular expression") != -1:
                raise IncorrectSyntax("Unsupported question type: regular expression")
            if question_announcement.find("reflective single choice") != -1:
                raise IncorrectSyntax("Unsupported question type: reflective single choice")
            if question_announcement.find("reflective multiple choice") != -1:
                raise IncorrectSyntax("Unsupported question type: reflective multiple choice")
            if question_announcement.find("reflective text answer") != -1:
                raise IncorrectSyntax("Unsupported question type: reflective text answer")

            #Adding other options
            if question_announcement.find("shuffle") != -1:
                options[1] = shuffle.SHUFFLE
            if question_announcement.find("no shuffle") != -1:
                options[1] = shuffle.NO_SHUFFLE        
            if question_announcement.find("partial credit") != -1:
                options[2] = partial_credit.PARTIAL_CREDIT
            if question_announcement.find("no partial credit") != -1:
                options[2] = partial_credit.NO_PARTIAL_CREDIT

            full_question = MathJax_border_fix(full_question)
            if options[0] ==  question_type.SINGLE_CHOICE:
                problem = create_multiple_choice_question(options, full_question)
            if options[0] ==  question_type.MULTIPLE_CHOICE:
                problem = create_checkbox_question(options, full_question)
            if options[0] ==  question_type.NUMERIC:
                problem = create_numeric_question(full_question)
            if options[0] ==  question_type.MATH_EXPRESSION:
                problem = create_math_expression_question(full_question)
            if options[0] ==  question_type.TEXT_MATCH:
                problem = create_text_match_question(full_question)

            problem_text = []
            problem.add_to_text(problem_text)
            problem_str = '\n'.join(problem_text)
            #print(problem_str + '\n')

            lib.add_problem("q" + str(q_i))
            problems.append(problem_str)
            q_i += 1
            i = end_of_question


        lib_text = []
        lib.add_to_text(lib_text)
        lib_str = '\n'.join(lib_text)
        #print(lib_str + '\n')
    except IncorrectSyntax as ex:
        print("The conversion was stopped for the following reasons:\n", ex.args[0])
        lib_str = ""
        problems = []
    except xml_elements.IncorrectSyntax as ex:
        print("?????? ?????????? ???????????? ??????????????????????, ???????????????????? ?????? ??????!\n", ex.args[0])
        lib_str = ""
        problems = []


    return lib_str, problems

        
